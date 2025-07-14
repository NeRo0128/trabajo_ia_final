# generar_informe.py
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

#* 1. Crear documento y portada
doc = Document()

# Título centrado, grande y en negrita
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Predicción de Complicaciones Clínicas Posteriores a Cirugías de Cadera mediante Aprendizaje Automático')
run.bold = True
run.font.size = Pt(22)
title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Añadir espacio vertical
doc.add_paragraph("\n" * 10)

# Tabla invisible para alinear al pie de página
footer_table = doc.add_table(rows=1, cols=1)
footer_table.allow_autofit = False
footer_cell = footer_table.cell(0, 0)

footer_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM  # <- Esta línea corregida

p1 = footer_cell.paragraphs[0]
p1.add_run("Autor: Rodolfo Gilberto Castillo Vega\n")
p1.add_run("Institución: Universidad Central Marta Abreu de las Villas\n")
p1.add_run("Carrera: 3er año de Licenciatura en Ciencias de la Computación")
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

#* 2. Introducción
doc.add_heading('Introducción', level=1)
doc.add_paragraph(
    "Las fracturas de cadera representan uno de los eventos clínicos más graves y frecuentes entre la población adulta mayor, provocando consecuencias significativas en términos de morbilidad, mortalidad y carga económica para los sistemas de salud. Debido al envejecimiento progresivo de la población, la incidencia de este tipo de fracturas ha aumentado considerablemente en las últimas décadas, y se espera que esta tendencia continúe en el futuro." +
    "Una vez que ocurre una fractura de cadera, el tratamiento más común incluye una intervención quirúrgica seguida de un periodo de recuperación y rehabilitación. Sin embargo, no todos los pacientes responden de la misma manera: algunos evolucionan favorablemente, mientras que otros presentan complicaciones que pueden afectar su calidad de vida, prolongar la estancia hospitalaria o incluso comprometer su supervivencia."+
    "En este contexto, surge la necesidad de aplicar herramientas de análisis de datos e inteligencia artificial para predecir con anticipación qué pacientes tienen mayor riesgo de desarrollar complicaciones postoperatorias. Esta información permitiría al personal médico tomar decisiones más informadas, personalizar los tratamientos, anticipar cuidados especiales y optimizar los recursos disponibles."+
    "El presente proyecto se enfoca en la construcción y evaluación de modelos de aprendizaje automático (machine learning) capaces de predecir complicaciones tras una operación de cadera. A partir de un conjunto de datos clínicos reales, se desarrolló un pipeline de procesamiento que incluye limpieza, codificación, escalado, entrenamiento de modelos, ajuste de hiperparámetros y evaluación con métricas estándares. Se emplearon tres modelos de clasificación ampliamente usados: Random Forest, K-Nearest Neighbors (KNN) y una red neuronal del tipo perceptrón multicapa (MLP)."+
    "El objetivo general de este trabajo es analizar el rendimiento de estos modelos y determinar cuál ofrece los mejores resultados en cuanto a precisión y capacidad de generalización. Además, se busca interpretar los factores que más influyen en la predicción de complicaciones y ofrecer recomendaciones para futuros estudios y aplicaciones clínicas."
)
doc.add_page_break()

#* 3. Métodos de Preprocesamiento
doc.add_heading('Métodos de Preprocesamiento', level=1)
# Descripción y justificación de cada paso aplicado en el script de preprocesamiento
methods = [
    ('Carga de datos', 'Importación del archivo Excel a un DataFrame de pandas para facilitar la manipulación estructurada de las variables.'),
    ('Manejo de valores nulos', 'Relleno de valores faltantes en la columna `tiempo_quirurgico` con la etiqueta "no_definido" para evitar pérdida de registros y marcar explícitamente la ausencia de información.'),
    ('Eliminación de columnas irrelevantes', 'Descartar variables sin aporte al modelo (p. ej., `no`, `evolucion_al_mes`, etc.) para reducir dimensionalidad y mejorar rendimiento.'),
    ('Codificación de variables categóricas', 'Transformación de variables nominales en valores numéricos mediante `LabelEncoder` para compatibilidad con los algoritmos de scikit-learn.'),
    ('Normalización de variables numéricas', 'Escalado de características con `StandardScaler` (media 0, desviación típica 1) para garantizar convergencia y estabilidad en modelos sensibles a la escala.'),
    ('Definición de X e y', 'Separar las variables predictoras (`X`) de la variable objetivo (`y = complicaciones`) para estructurar el entrenamiento.'),
    ('División entrenamiento-prueba', 'Uso de `train_test_split` (80 %–20 %) con `stratify=y` y `random_state=42` para una evaluación robusta manteniendo la proporción de clases.'),
    ('Persistencia de artefactos', 'Guardado de conjuntos preprocesados y objetos de transformación (`LabelEncoder`, `StandardScaler`) con `joblib` para reproducibilidad y despliegue.')
]
for title, desc in methods:
    doc.add_paragraph(f"• {title}: {desc}")

doc.add_page_break()

#* 4. EDA
df = pd.read_excel('data/processed/base_modificada.xlsx')

# 4.1 Distribución de complicaciones
plt.figure(figsize=(6,4))
df['complicaciones'].value_counts().plot(kind='bar', color='skyblue')
plt.title('Distribución de Complicaciones')
plt.xlabel('Complicaciones')
plt.ylabel('Frecuencia')
plt.tight_layout()
plt.savefig('reports/img/eda_complicaciones.png')
doc.add_heading('Análisis Exploratorio de Datos (EDA)', level=1)
doc.add_picture('reports/img/eda_complicaciones.png', width=Inches(5))
doc.add_paragraph('Figura 1: Distribución de la variable objetivo “complicaciones”.')

# 4.2 Matriz de correlación
corr = df[['edad', 't_max_(oc)', 't_min_(oc)', 't_med_(oc)', 'hr_med_(%)', 'pe_med_(hpa)']].corr()
plt.figure(figsize=(6,6))
sns.heatmap(corr, annot=True, cmap='viridis')
plt.title('Matriz de Correlación')
plt.tight_layout()
plt.savefig('reports/img/eda_correlacion.png')
doc.add_picture('reports/img/eda_correlacion.png', width=Inches(5))
doc.add_paragraph('Figura 2: Matriz de correlación de variables numéricas.')
doc.add_page_break()

#* 5. Modelos y resultados
doc.add_heading('Modelos y Búsqueda de Hiperparámetros', level=1)
# Cargar CSVs de resultados
rf = pd.read_csv('data/results/rf_results.csv')
knn = pd.read_csv('data/results/knn_results.csv')
mlp = pd.read_csv('data/results/mlp_results.csv')

# 5.1 Random Forest
doc.add_heading('Random Forest', level=2)
tbl_rf = doc.add_table(rows=1, cols=4)
hdr = tbl_rf.rows[0].cells
for i, col in enumerate(['n_estimators','max_depth','min_samples_split','accuracy']):
    hdr[i].text = col
for _, r in rf.iterrows():
    row = tbl_rf.add_row().cells
    row[0].text = str(r['n_estimators'])
    row[1].text = str(r['max_depth'])
    row[2].text = str(r['min_samples_split'])
    row[3].text = f"{r['accuracy']:.3f}"
doc.add_page_break()

# 5.2 KNN
doc.add_heading('K-Nearest Neighbors', level=2)
tbl_knn = doc.add_table(rows=1, cols=3)
hdr = tbl_knn.rows[0].cells
for i, col in enumerate(['n_neighbors','weights','accuracy']):
    hdr[i].text = col
for _, r in knn.iterrows():
    row = tbl_knn.add_row().cells
    row[0].text = str(r['n_neighbors'])
    row[1].text = r['weights']
    row[2].text = f"{r['accuracy']:.3f}"
doc.add_page_break()

# 5.3 MLPClassifier
doc.add_heading('Red Neuronal (MLPClassifier)', level=2)
tbl_mlp = doc.add_table(rows=1, cols=6)
hdr = tbl_mlp.rows[0].cells
for i, col in enumerate(['hidden_layer_sizes','alpha','learning_rate_init','max_iter','early_stopping','accuracy']):
    hdr[i].text = col
for _, r in mlp.iterrows():
    row = tbl_mlp.add_row().cells
    row[0].text = str(tuple(r['hidden_layer_sizes']).tolist() if isinstance(r['hidden_layer_sizes'], (list,tuple)) else r['hidden_layer_sizes'])
    row[1].text = str(r['alpha'])
    row[2].text = str(r['learning_rate_init'])
    row[3].text = str(int(r['max_iter']))
    row[4].text = str(r['early_stopping'])
    row[5].text = f"{r['accuracy']:.3f}"



#* 6. Selección de mejores parámetros

doc.add_heading('Selección de Mejores Parámetros', level=1)

# Identificar los mejores registros
best_rf = rf.loc[rf['accuracy'].idxmax()]
best_knn = knn.loc[knn['accuracy'].idxmax()]
best_mlp = mlp.loc[mlp['accuracy'].idxmax()]

# Agregar análisis argumentado
doc.add_paragraph(
    "Tras realizar la búsqueda de hiperparámetros mediante exploración en cuadrícula (grid search), "
    "se identificaron los conjuntos óptimos para cada uno de los modelos evaluados. Esta selección se "
    "basó en el valor más alto de la métrica de exactitud (accuracy) obtenida sobre el conjunto de prueba."
)

# Random Forest
doc.add_paragraph(
    f"Para el modelo Random Forest, la configuración con mejor desempeño fue: "
    f"n_estimators = {best_rf['n_estimators']}, max_depth = {best_rf['max_depth']}, "
    f"min_samples_split = {best_rf['min_samples_split']}. "
    f"Esta combinación logró una accuracy de {best_rf['accuracy']:.3f}, lo cual indica que el modelo fue capaz "
    "de aprender correctamente patrones relevantes incluso con un número moderado de árboles, sin necesidad de limitar su profundidad."
)

# KNN
doc.add_paragraph(
    f"En el caso de K-Nearest Neighbors (KNN), el mejor resultado se obtuvo con: "
    f"n_neighbors = {best_knn['n_neighbors']}, weights = {best_knn['weights']}. "
    f"Este modelo también alcanzó una accuracy de {best_knn['accuracy']:.3f}, lo que sugiere que, en este conjunto de datos, "
    "la clase de una muestra puede inferirse de forma confiable a partir de sus 3 vecinos más cercanos sin ponderación adicional por distancia."
)

# MLPClassifier
doc.add_paragraph(
    f"Para el clasificador de red neuronal (MLPClassifier), la arquitectura más efectiva fue: "
    f"hidden_layer_sizes = {best_mlp['hidden_layer_sizes']}, alpha = {best_mlp['alpha']}, "
    f"learning_rate_init = {best_mlp['learning_rate_init']}, max_iter = {int(best_mlp['max_iter'])}, "
    f"early_stopping = {best_mlp['early_stopping']}. Este modelo alcanzó igualmente una accuracy de {best_mlp['accuracy']:.3f}, "
    "indicando que una red de una sola capa oculta con 100 neuronas, junto con una regularización y tasa de aprendizaje moderadas, "
    "fue suficiente para capturar las relaciones no lineales presentes en los datos."
)

doc.add_paragraph(
    "Aunque los tres modelos alcanzaron el mismo nivel de exactitud, se debe tener en cuenta el tiempo de entrenamiento, "
    "la interpretabilidad y la robustez de cada uno. En este sentido, Random Forest resulta una opción balanceada, ya que combina "
    "buen desempeño con interpretabilidad (importancia de variables) y menor riesgo de sobreajuste en comparación con MLP."
)

#* 7. Conclusiones

doc.add_page_break()
doc.add_heading('Conclusiones y Trabajo Futuro', level=1)
doc.add_paragraph(
    "Random Forest y KNN alcanzaron 0.917 de accuracy, siendo modelos robustos y eficientes. "
    "MLP presentó mismos resultados pero con mayor tiempo de entrenamiento. "
    "Se recomienda Random Forest para implementación clínica. "
    "Futuros trabajos: aumentar tamaño de muestra, explorar ensambles y validación externa."
)

# Guardar
doc.save('reports/informe_completo.docx')
print("Informe generado: informe_completo.docx")
